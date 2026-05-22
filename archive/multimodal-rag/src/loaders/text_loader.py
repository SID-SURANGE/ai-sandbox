# Standard library imports
import os
import uuid

# third-party imports
import fitz
from docx import Document
from pathlib import Path
import logging

# local imports
from utils.config import collection, text_splitter, text_model

logger = logging.getLogger(__name__)

def generate_id(prefix: str, filename: str, chunk_num: int) -> str:
    """Generate a unique ID for a chunk."""
    return f"{prefix}_{Path(filename).stem}_{chunk_num}_{uuid.uuid4().hex[:8]}"

def process_text(file_path):
    """Process a text file, generate embeddings, and store in ChromaDB."""
    try:
        print(f"Processing text file: {file_path}")
        
        # Read the text content
        with open(file_path, "r", encoding="utf-8") as f:
            text_content = f.read()

        # Split text into chunks
        chunks = text_splitter.split_text(text_content)
        
        if not chunks:
            print(f"No content found in file: {file_path}")
            return False

        # Lists to store chunk data
        embeddings = []
        metadatas = []
        documents = []
        ids = []

        # Process each chunk
        for chunk_num, chunk in enumerate(chunks):
            # Generate embedding using text model
            chunk_embedding = text_model.encode(chunk)
            
            # Generate unique ID
            chunk_id = generate_id("txt", file_path, chunk_num)
            
            # Prepare metadata according to unified schema
            metadata = {
                "content_type": "text",
                "file_path": str(file_path),
                "filename": Path(file_path).name,
                "description": chunk,  # Store full chunk as description
                "keywords": ""  # Empty string for now, can be populated with key terms if needed
            }
            
            documents.append(str(file_path))  # Convert Path to string
            embeddings.append(chunk_embedding.tolist())
            metadatas.append(metadata)
            ids.append(chunk_id)

        # Store in ChromaDB
        collection.add(
            documents=[str(file_path) for _ in range(len(chunks))],  # Convert Path to string
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids
        )
        
        print(f"Text file '{file_path}' processed and added to ChromaDB with {len(chunks)} chunks.")
        return True

    except Exception as e:
        logger.error(f"Error processing text file {file_path}: {str(e)}")
        print(f"Failed to process text file {file_path}: {str(e)}")
        return False


def process_pdf(file_path):
    """Process a PDF file, extract text, generate embeddings, and store in ChromaDB."""
    try:
        print(f"Processing PDF file: {file_path}")
        
        # Validate file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
            
        # Extract content from PDF
        try:
            doc = fitz.open(file_path)
            text_list = []
            
            # Process each page
            for page in doc:
                text = page.get_text()
                if text.strip():
                    text_list.append(text)
                    
            if not text_list:
                raise ValueError("PDF appears to be empty or contains no extractable content")
                
            # Join all text and split into chunks
            full_text = "\n".join(text_list)
            chunks = text_splitter.split_text(full_text) if text_list else []
                
        except Exception as e:
            raise ValueError(f"Failed to extract content from PDF: {str(e)}")
        finally:
            doc.close()
            
        # Process and store content
        if chunks:
            documents = []
            embeddings = []
            metadatas = []
            ids = []
            
            for chunk_num, chunk in enumerate(chunks):
                # Generate embedding using text model
                embedding = text_model.encode(chunk)
                
                # Generate unique ID
                chunk_id = generate_id("pdf", file_path, chunk_num)
                
                # Prepare metadata according to unified schema
                metadata = {
                    "content_type": "pdf",
                    "file_path": str(file_path),
                    "filename": Path(file_path).name,
                    "description": chunk,  # Store full chunk as description
                    "keywords": ""  # Empty string for now, can be populated with key terms if needed
                }
                
                documents.append(str(file_path))  # Convert Path to string
                embeddings.append(embedding.tolist())
                metadatas.append(metadata)
                ids.append(chunk_id)
            
            # Store in ChromaDB
            collection.add(
                documents=[str(file_path) for _ in range(len(chunks))],  # Convert Path to string
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            print(f"PDF file '{file_path}' processed and added to ChromaDB with {len(chunks)} chunks.")
            return True
                
        else:
            print(f"No valid content to add from file: {file_path}")
            return False
            
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {str(e)}")
        print(f"Failed to process PDF {file_path}: {str(e)}")
        return False

def process_docx(file_path):
    """Process a DOCX file, extract text, generate embeddings, and store in ChromaDB."""
    try:
        print(f"Processing DOCX file: {file_path}")
        
        # Validate file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        # Load and extract text from DOCX
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    full_text.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            # Join all text
            document_text = "\n".join(full_text)
            
            if not document_text.strip():
                raise ValueError("Document appears to be empty")
                
            # Split into chunks
            chunks = text_splitter.split_text(document_text)
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        # Process chunks and store in ChromaDB
        if chunks:
            documents = []
            embeddings = []
            metadatas = []
            ids = []
            
            for chunk_num, chunk in enumerate(chunks):
                # Generate embedding using text model
                embedding = text_model.encode(chunk)
                
                # Generate unique ID
                chunk_id = generate_id("docx", file_path, chunk_num)
                
                # Prepare metadata according to unified schema
                metadata = {
                    "content_type": "docx",
                    "file_path": str(file_path),
                    "filename": Path(file_path).name,
                    "description": chunk,  # Store full chunk as description
                    "keywords": ""  # Empty string for now, can be populated with key terms if needed
                }
                
                documents.append(str(file_path))  # Convert Path to string
                embeddings.append(embedding.tolist())
                metadatas.append(metadata)
                ids.append(chunk_id)
            
            # Store in ChromaDB
            collection.add(
                documents=[str(file_path) for _ in range(len(chunks))],  # Convert Path to string
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            print(f"DOCX file '{file_path}' processed and added to ChromaDB with {len(chunks)} chunks.")
            return True
            
        else:
            print(f"No valid content to add from file: {file_path}")
            return False
            
    except Exception as e:
        logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        print(f"Failed to process DOCX {file_path}: {str(e)}")
        return False
